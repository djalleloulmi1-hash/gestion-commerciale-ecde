from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in kwargs:
        element = OxmlElement(f'w:{border_name}')
        element.set(qn('w:val'), kwargs[border_name])
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcPr.append(element)

def create_readme_docx():
    doc = Document()
    
    # --- HEADER & LOGO ---
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_path = "logo_entete.png"
    if os.path.exists(logo_path):
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(6.5)) # Slightly wider for full header effect
    else:
        paragraph.add_run("GROUPE GICA - ECDE").bold = True

    # --- TITLE ---
    doc.add_paragraph("\n")
    title = doc.add_heading("GESTION COMMERCIALE GICA (ECDE)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("MANUEL D'INSTALLATION ET D'UTILISATION")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True

    doc.add_paragraph("\n")

    # --- IMPORTANT WARNING ---
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_alert = p.add_run("IMPORTANT : A LIRE IMPERATIVEMENT AVANT LA PREMIERE UTILISATION")
    run_alert.bold = True
    run_alert.font.color.rgb = RGBColor(255, 0, 0) # Red
    
    p.add_run("\n\nCe logiciel est configuré pour fonctionner dans un répertoire spécifique pour garantir la stabilité de la base de données et des sauvegardes.")
    
    set_cell_border(cell, top="single", bottom="single", left="single", right="single")
    
    doc.add_paragraph("\n")

    # --- 1. INSTALLATION ---
    doc.add_heading("📥 INSTALLATION RAPIDE", level=1)
    doc.add_paragraph("Pour que le logiciel fonctionne correctement, veuillez suivre scrupuleusement ces étapes lors de la première installation :")

    doc.add_heading("1. Emplacement du Dossier (CRITIQUE)", level=2)
    p = doc.add_paragraph()
    runner = p.add_run("Le logiciel DOIT être installé directement à la racine du disque C:.")
    runner.bold = True
    
    doc.add_paragraph("1. Copiez le dossier complet 'GICA_PROJET'.", style="List Number")
    doc.add_paragraph("2. Collez-le dans C:\\ (Disque Local C).", style="List Number")
    
    p = doc.add_paragraph("3. Le chemin final doit être exactement : ")
    p.style = "List Number"
    p.add_run("C:\\GICA_PROJET").bold = True
    
    doc.add_paragraph("Ne pas installer dans 'Mes Documents', 'Bureau', ou 'Program Files'.", style="List Bullet")

    doc.add_heading("2. Lancement du Logiciel", level=2)
    doc.add_paragraph("Pour un accès facile :")
    doc.add_paragraph("1. Allez dans C:\\GICA_PROJET.", style="List Number")
    doc.add_paragraph("2. Faites un clic droit sur 'Lancer gestion.bat'.", style="List Number")
    doc.add_paragraph("3. Choisissez 'Envoyer vers' > 'Bureau (créer un raccourci)'.", style="List Number")

    # --- 2. PREMIERE CONNEXION ---
    doc.add_heading("🔑 PREMIÈRE CONNEXION", level=1)
    doc.add_paragraph("Identifiants Administrateur par défaut :")
    doc.add_paragraph("Identifiant : admin", style="List Bullet")
    doc.add_paragraph("Mot de passe : admin", style="List Bullet")
    p = doc.add_paragraph("Conseil : Changez ce mot de passe via le menu Configuration -> Utilisateurs.")
    p.runs[0].italic = True

    # --- 3. GUIDE UTILISATION ---
    doc.add_heading("🚀 GUIDE D'UTILISATION DÉTAILLÉ", level=1)
    
    doc.add_heading("1. Tableau de Bord (Dashboard)", level=2)
    doc.add_paragraph("Dès la connexion, vue synthétique de l'activité : Chiffre d'Affaires, État du Stock 104, et Tonnage vendu.")

    doc.add_heading("2. Gestion des Clients", level=2)
    doc.add_paragraph("Gérez votre base de données partenaires :")
    doc.add_paragraph("- Seuil de Crédit : Bloque la facturation si le client dépasse son plafond autorisé.", style="List Bullet")
    doc.add_paragraph("- Solde en temps réel : Rouge = Créditeur (Doit de l'argent), Orange = Débiteur.", style="List Bullet")

    doc.add_heading("3. Gestion des Stocks (Réceptions)", level=2)
    doc.add_paragraph("Enregistrement des entrées de marchandises (Achats usine).")
    doc.add_paragraph("- Saisie simple : Fournisseur, Date, Quantités (Annoncée vs Reçue).", style="List Bullet")
    doc.add_paragraph("- Sécurité : La suppression est virtuelle ('Soft Delete') pour garder une trace comptable.", style="List Bullet")

    doc.add_heading("4. Facturation (Ventes)", level=2)
    doc.add_paragraph("Le cœur du système :")
    doc.add_paragraph("1. Sélection Client (Solde et Seuil affichés).", style="List Number")
    doc.add_paragraph("2. Ajout Produits (Contrôle de stock immédiat).", style="List Number")
    doc.add_paragraph("3. Validation : Génération PDF, Débit Stock, Créance Client.", style="List Number")
    doc.add_paragraph("Note : Pour annuler, utilisez le bouton 'Créer Avoir' qui réintègre automatiquement le stock.", style="List Bullet")

    doc.add_heading("5. Paiements", level=2)
    doc.add_paragraph("Encaissez Espèces, Chèques ou Virements. Les paiements diminuent instantanément le solde client.")

    # --- 4. ETATS GENERES ---
    doc.add_heading("📊 ÉTATS ET RAPPORTS", level=1)
    
    reports = [
        ("Factures de Vente", "Document commercial normalisé."),
        ("Bons de Livraison", "Justificatif de sortie de stock."),
        ("Bordereaux de Versement", "Pour la remise de chèques en banque."),
        ("État 104 (CA)", "Tableau fiscal par client (NIF, Art Impo)."),
        ("État des Stocks (Inventaire)", "Quantités physiques disponibles."),
        ("Mouvements Stock (Valorisé)", "Historique Entrées/Sorties avec valeurs financières."),
        ("Situation Journalière", "Résumé Activité, Ventes, CA, Trésorerie."),
        ("Consommation Globale", "Rapport détaillé par période.")
    ]
    
    for title_text, desc_text in reports:
        p = doc.add_paragraph()
        runner_title = p.add_run(f"• {title_text} : ")
        runner_title.bold = True
        runner_desc = p.add_run(desc_text)

    # --- 5. MAINTENANCE ---
    doc.add_heading("🛠️ MAINTENANCE ET SÉCURITÉ", level=1)
    
    p = doc.add_paragraph("Sauvegardes Automatiques")
    p.style = "Heading 2"
    doc.add_paragraph("Copie de sécurité créée à chaque fermeture dans C:\\GICA_PROJET\\Backups.")
    
    p = doc.add_paragraph("Auto-Réparation")
    p.style = "Heading 2"
    doc.add_paragraph("Le logiciel vérifie et répare sa propre base de données au démarrage si nécessaire.")

    # --- FOOTER ---
    doc.add_paragraph("\n" * 2)
    line = doc.add_paragraph("_" * 50)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Développé par Mr OULMI ABDELDJALLIL").bold = True
    p.add_run("\nN° Tel: 0554 15 57 37")
    p.add_run("\nE-mail: djalleloulmi1@gmail.com")

    # Save
    filename = "READ ME (Lisez Moi) Gestion Commerciale GICA.docx"
    doc.save(filename)
    print(f"Document créé : {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_readme_docx()
