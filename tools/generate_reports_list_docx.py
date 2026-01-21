from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_reports_list():
    doc = Document()
    
    # Title
    title = doc.add_heading("ETATS GENERES PAR LE PROGRAMME", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    
    intro = doc.add_paragraph("Le logiciel assure l'édition automatique de l'ensemble des documents commerciaux et fiscaux nécessaires à la bonne gestion de l'unité :")
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # List of Reports
    reports = [
        ("Factures de Vente", "Document commercial normalisé incluant le calcul automatique des montants HT, TVA, et Droit de Timbre. Chaque facture est numérotée séquentiellement."),
        
        ("Bons de Livraison", "Générés systématiquement avec chaque facture, ils servent de justificatif de sortie de stock et de transport."),
        
        ("Bordereaux de Versement", "État préparatoire pour la banque, regroupant les chèques et virements encaissés une journée donnée."),
        
        ("État 104 (Chiffre d'Affaires)", "Tableau récapitulatif fiscal du Chiffre d'Affaires par client. Il détaille l'Article d'Imposition, le NIF, et les montants cumulés sur la période sélectionnée."),
        
        ("État des Stocks (Inventaire)", "Situation en temps réel des quantités physiques disponibles pour chaque produit (Ciment, Sacs, Vrac)."),
        
        ("État des Mouvements de Stocks (Valorisé)", "Rapport détaillé traçant l'historique des Entrées (Réceptions usine) et des Sorties (Ventes), avec une valorisation financière des mouvements."),
        
        ("Situation Journalière & Globale", "Tableau de bord imprimable résumant l'activité de la journée : Quantités vendues, CA réalisé, et situation de la trésorerie."),
        
        ("Relevé de Compte Client (Solde)", "Fiche détaillée par client montrant l'historique chronologique de ses opérations (Factures vs Paiements) et son solde actuel (Débit/Crédit).")
    ]
    
    for title_text, desc_text in reports:
        p = doc.add_paragraph()
        runner_title = p.add_run(f"• {title_text} : ")
        runner_title.bold = True
        runner_title.font.size = Pt(12)
        
        runner_desc = p.add_run(desc_text)
        runner_desc.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    # Save
    filename = "Etats_Generes_Par_Le_Programme.docx"
    doc.save(filename)
    print(f"Document généré : {os.path.abspath(filename)}")

if __name__ == "__main__":
    generate_reports_list()
