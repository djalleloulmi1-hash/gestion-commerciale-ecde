from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in kwargs:
        element = OxmlElement(f'w:{border_name}')
        element.set(qn('w:val'), kwargs[border_name])
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcPr.append(element)

def create_detailed_manual():
    doc = Document()
    
    # 0. Header with Logo
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_path = "logo_entete.png"
    if os.path.exists(logo_path):
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(6))
    else:
        paragraph.add_run("GROUPE GICA - ECDE").bold = True

    # 1. Title Page
    doc.add_paragraph("\n" * 2)
    title = doc.add_heading("LOGICIEL DE GESTION COMMERCIALE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Guide Utilisateur Complet")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.name = 'Calibri Light'
    
    doc.add_paragraph("\n" * 3)
    
    # Version Info Frame
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("VERSION 1.0").bold = True
    p.add_run("\n\nCette version constitue la première release officielle du logiciel.\nElle couvre l'essentiel des opérations de gestion commerciale.\nLes futures mises à jour intégreront des améliorations basées\nsur vos retours et les besoins exprimés par la direction.").italic = True
    
    set_cell_border(cell, top="single", bottom="single", left="single", right="single")
    
    doc.add_page_break()

    # 2. Table of Contents (Simulated)
    doc.add_heading("Sommaire", level=1)
    doc.add_paragraph("1. Présentation Générale")
    doc.add_paragraph("2. Installation et Lancement")
    doc.add_paragraph("3. Prise en Main (Première Connexion)")
    doc.add_paragraph("4. Fonctionnalités Détaillées")
    doc.add_paragraph("5. États Générés par le Programme")
    doc.add_paragraph("6. Maintenance et Sécurité")
    doc.add_paragraph("7. Support Technique")
    doc.add_paragraph("\n")

    # 3. Content
    
    # 3.1 Présentation
    doc.add_heading("1. Présentation Générale", level=1)
    p = doc.add_paragraph("Ce logiciel a été conçu sur mesure pour l'ECDE (Groupe GICA) afin d'informatiser et de sécuriser la gestion commerciale de l'unité de distribution. Il centralise toutes les opérations : des réceptions de ciment aux ventes, en passant par la gestion des clients et des stocks.")
    
    # 3.2 Installation
    doc.add_heading("2. Installation et Lancement", level=1)
    doc.add_paragraph("L'installation a été simplifiée au maximum pour éviter toute erreur technique.")
    
    doc.add_heading("2.1. Emplacement Impératif", level=2)
    p = doc.add_paragraph()
    runner = p.add_run("Le dossier 'GESTION ECDE 2026' doit être copié directement à la racine du disque C:")
    runner.bold = True
    runner.font.color.rgb = RGBColor(200, 0, 0) # Dark Red
    
    doc.add_paragraph("Chemin correct : C:\\GICA_PROJET", style="List Bullet")
    doc.add_paragraph("Ne pas placer sur le Bureau ou dans 'Mes Documents'", style="List Bullet")
    
    doc.add_heading("2.2. Lancement", level=2)
    doc.add_paragraph("Dans le dossier, double-cliquez simplement sur le fichier :")
    p = doc.add_paragraph("Lancer Application.bat")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # 3.3 Connexion
    doc.add_heading("3. Prise en Main", level=1)
    doc.add_paragraph("Au lancement, une fenêtre d'authentification apparaît.")
    doc.add_paragraph("Identifiants Administrateur par défaut :", style="List Bullet")
    doc.add_paragraph("Utilisateur : admin", style="List Bullet")
    doc.add_paragraph("Mot de passe : admin", style="List Bullet")
    
    # 3.4 Fonctionnalités
    doc.add_heading("4. Fonctionnalités Détaillées", level=1)
    
    doc.add_heading("4.1. Tableau de Bord (Dashboard)", level=2)
    doc.add_paragraph("Dès la connexion, vous avez une vue synthétique de l'activité :")
    doc.add_paragraph("- Chiffre d'Affaires journalier et mensuel.")
    doc.add_paragraph("- État du Stock 104 en temps réel.")
    doc.add_paragraph("- Nombre de camions et tonnage vendu.")
    
    doc.add_heading("4.2. Gestion des Clients", level=2)
    doc.add_paragraph("L'onglet 'Clients' permet de créer et modifier les fiches clients. Chaque client possède :")
    doc.add_paragraph("- Un Seuil de Crédit : Si dépassé, le logiciel bloque la facturation (Sécurité).", style="List Bullet")
    doc.add_paragraph("- Un Solde en temps réel : Calculé automatiquement (Report N-1 + Ventes - Paiements).", style="List Bullet")
    
    doc.add_heading("4.3. Facturation", level=2)
    doc.add_paragraph("Le module de facturation est le cœur du système :")
    doc.add_paragraph("1. Sélectionnez le client (son solde s'affiche).")
    doc.add_paragraph("2. Ajoutez les produits (le stock est vérifié instantanément).")
    doc.add_paragraph("3. Saisissez les infos transport (Matricule, Chauffeur).")
    doc.add_paragraph("4. Validez : La facture s'imprime et le stock est débité.")
    
    doc.add_heading("4.4. Avoirs et Annulations", level=2)
    doc.add_paragraph("Pour annuler une facture ou faire un retour :")
    doc.add_paragraph("- Sélectionnez la facture dans la liste.")
    doc.add_paragraph("- Cliquez sur 'Créer Avoir'.")
    doc.add_paragraph("- Le stock est automatiquement réintégré.")

    # 3.5 Etats Generes (NEW SECTION)
    doc.add_heading("5. ETATS GENERES PAR LE PROGRAMME", level=1)
    doc.add_paragraph("Le logiciel assure l'édition automatique de l'ensemble des documents commerciaux et fiscaux nécessaires à la bonne gestion de l'unité :")
    
    reports = [
        ("Factures de Vente", "Document commercial normalisé incluant le calcul automatique des montants HT, TVA, et Droit de Timbre. Chaque facture est numérotée séquentiellement."),
        ("Bons de Livraison", "Générés systématiquement avec chaque facture, ils servent de justificatif de sortie de stock et de transport."),
        ("Bordereaux de Versement", "État préparatoire pour la banque, regroupant les chèques et virements encaissés une journée donnée."),
        ("État 104 (Chiffre d'Affaires)", "Tableau récapitulatif fiscal du Chiffre d'Affaires par client. Il détaille l'Article d'Imposition, le NIF, et les montants cumulés sur la période sélectionnée."),
        ("État des Stocks (Inventaire)", "Situation en temps réel des quantités physiques disponibles pour chaque produit (Ciment, Sacs, Vrac)."),
        ("État des Mouvements de Stocks (Valorisé)", "Rapport détaillé traçant l'historique des Entrées (Réceptions usine) et des Sorties (Ventes), avec une valorisation financière des mouvements."),
        ("Situation Journalière & Globale", "Tableau de bord imprimable résumant l'activité de la journée : Quantités vendues, CA réalisé, et situation de la trésorerie."),
        ("Relevé de Compte Client (Solde)", "Fiche détaillée par client montrant l'historique chronologique de ses opérations (Factures vs Paiements) et son solde actuel (Débit/Crédit).")
    ]
    
    for title_text, desc_text in reports:
        p = doc.add_paragraph()
        runner_title = p.add_run(f"• {title_text} : ")
        runner_title.bold = True
        runner_desc = p.add_run(desc_text)
    
    # 3.6 Maintenance
    doc.add_heading("6. Maintenance et Sécurité", level=1)
    doc.add_paragraph("Le logiciel est équipé de systèmes de sécurité autonomes :")
    
    p = doc.add_paragraph("Sauvegardes Automatiques")
    p.style = "Heading 2"
    doc.add_paragraph("À chaque fermeture, une copie de la base de données est créée dans le dossier 'Backups'. En cas de problème informatique, aucune donnée n'est perdue.")
    
    p = doc.add_paragraph("Auto-Réparation (Self-Healing)")
    p.style = "Heading 2"
    doc.add_paragraph("Au démarrage, le logiciel vérifie l'intégrité de ses fichiers. S'il détecte une anomalie dans la base de données, il tente de la réparer automatiquement.")

    # 4. Support Footer
    doc.add_paragraph("\n" * 3)
    line = doc.add_paragraph("_" * 50)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Développé par Mr OULMI ABDELDJALLIL").bold = True
    p.add_run("\nN° Tel: 0554 15 57 37")
    p.add_run("\nE-mail: djalleloulmi1@gmail.com")
    
    filename = "Guide_Utilisateur_Complet_v1.0.docx"
    doc.save(filename)
    print(f"Document généré : {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_detailed_manual()
