from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_manual():
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # HEADER / TITLE
    title = doc.add_heading('GESTION COMMERCIALE GICA (ECDE)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('MANUEL D\'INSTALLATION ET D\'UTILISATION')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph().add_run().add_break()

    # IMPORTANT ALERT
    p = doc.add_paragraph()
    run = p.add_run("IMPORTANT : A LIRE IMPERATIVEMENT AVANT LA PREMIERE UTILISATION")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph("Ce logiciel est configuré pour fonctionner dans un répertoire spécifique pour garantir la stabilité de la base de données et des sauvegardes.")

    doc.add_paragraph().add_run().add_break()

    # 1. INSTALLATION
    doc.add_heading('1. INSTALLATION RAPIDE', level=1)
    
    doc.add_heading('1.1. Emplacement du Dossier', level=2)
    p = doc.add_paragraph("Le logiciel DOIT être installé directement à la racine du disque C:.")
    
    items = [
        "Copiez le dossier complet 'GICA_PROJET'.",
        "Collez-le dans C:\\ (Disque Local C).",
        "Le chemin final doit être : C:\\GICA_PROJET"
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')
        
    p = doc.add_paragraph()
    run = p.add_run("ATTENTION : Ne pas renommmer, ne pas mettre sur le Bureau.")
    run.bold = True
    
    doc.add_heading('1.2. Lancement du Logiciel', level=2)
    doc.add_paragraph("Pour faciliter l'utilisation quotidienne :")
    
    items = [
        "Allez dans le dossier C:\\GICA_PROJET.",
        "Trouvez le fichier 'Lancer gestion.bat'.",
        "Faites Clic Droit > Envoyer vers > Bureau (créer un raccourci).",
        "Lancez le logiciel depuis ce raccourci."
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')

    # 2. PREMIERE CONNEXION
    doc.add_heading('2. PREMIÈRE CONNEXION', level=1)
    doc.add_paragraph("Identifiants par défaut :")
    p = doc.add_paragraph()
    p.add_run("Identifiant : ").bold = True
    p.add_run("admin")
    p = doc.add_paragraph()
    p.add_run("Mot de passe : ").bold = True
    p.add_run("admin123")
    
    doc.add_paragraph("Il est recommandé de changer ce mot de passe via le menu Configuration -> Utilisateurs.").italic = True

    # 3. FONCTIONNALITES
    doc.add_heading('3. FONCTIONNALITÉS PRINCIPALES', level=1)
    
    doc.add_heading('3.1. Tableau de Bord', level=2)
    doc.add_paragraph("Vue d'ensemble du Chiffre d'Affaires, des Tonnes Vendues et de l'État du Stock en temps réel.")
    
    doc.add_heading('3.2. Ventes (Facturation)', level=2)
    doc.add_paragraph("Création intuitive de factures. Le système vérifie automatiquement le stock et le crédit client. Impression PDF automatique.")
    
    doc.add_heading('3.3. Stocks & Réceptions', level=2)
    doc.add_paragraph("Saisie des entrées (Achats) et traçabilité complète des mouvements.")
    
    doc.add_heading('3.4. Clients', level=2)
    doc.add_paragraph("Suivi des soldes en temps réel (Rouge = Créditeur, Orange = Débiteur).")
    
    doc.add_heading('3.5. Rapports', level=2)
    doc.add_paragraph("Génération des rapports : État des Ventes, Stock Valorisé, Consommation Globale.")

    # 4. MAINTENANCE
    doc.add_heading('4. MAINTENANCE ET SAUVEGARDE', level=1)
    doc.add_paragraph("Sauvegarde Automatique : Créée à chaque fermeture dans C:\\GICA_PROJET\\Backups.")

    doc.add_paragraph().add_run().add_break()
    
    # FOOTER / CONTACT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Développé par Mr: OULMI ABDELDJALLIL")
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("N° Téléphone : 0554 15 57 37")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("E-mail : djalleloulmi1@gmail.com")

    # Save
    pdf_path = "Manuel_Utilisateur_GICA.docx"
    doc.save(pdf_path)
    print(f"Manuel généré : {os.path.abspath(pdf_path)}")

if __name__ == "__main__":
    create_manual()
