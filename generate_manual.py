from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_manual():
    doc = Document()

    # Title
    title = doc.add_heading('Manuel Utilisateur', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Application de Gestion Commerciale - GICA')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Ce manuel décrit l'utilisation de l'application de Gestion Commerciale ECDE. "
        "Cette application permet de gérer les clients, les produits, les stocks, la facturation et les paiements."
    )

    # 2. Lancement
    doc.add_heading('2. Lancement de l\'application', level=1)
    doc.add_paragraph(
        "Pour lancer l'application :"
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Naviguez vers le dossier ").bold = True
    p.add_run("GestionCommerciale_GICA")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Double-cliquez sur le fichier ").bold = True
    p.add_run("GestionCommerciale_GICA.exe")
    
    doc.add_paragraph(
        "L'écran de connexion s'affichera."
    )

    # 3. Connexion
    doc.add_heading('3. Connexion', level=1)
    doc.add_paragraph(
        "Entrez vos identifiants pour accéder au système."
    )
    p = doc.add_paragraph()
    p.add_run("Utilisateur par défaut : ").bold = True
    p.add_run("admin")
    p = doc.add_paragraph()
    p.add_run("Mot de passe par défaut : ").bold = True
    p.add_run("admin123")

    # 4. Fonctionnalités
    doc.add_heading('4. Fonctionnalités Principales', level=1)
    
    doc.add_heading('4.1 Tableau de Bord', level=2)
    doc.add_paragraph(
        "La page d'accueil affiche un résumé de l'activité :\n"
        "- Chiffre d'Affaires du jour\n"
        "- Nombre de factures\n"
        "- Alertes de stock"
    )

    doc.add_heading('4.2 Clients', level=2)
    doc.add_paragraph(
        "Permet d'ajouter, modifier et consulter la liste des clients. "
        "Vous pouvez également suivre leur solde et leur historique."
    )

    doc.add_heading('4.3 Produits & Stocks', level=2)
    doc.add_paragraph(
        "Gérez votre catalogue produits et suivez les mouvements de stock (Réceptions, Ventes). "
        "Le stock est mis à jour automatiquement lors de la validation des factures."
    )

    doc.add_heading('4.4 Facturation', level=2)
    doc.add_paragraph(
        "Créez des factures proforma ou définitives. "
        "Une fois validée, la facture génère un PDF prêt à être imprimé."
    )

    # 5. Support
    doc.add_heading('5. Support Technique', level=1)
    doc.add_paragraph(
        "Pour toute question ou problème technique, veuillez contacter l'administrateur système."
    )

    # Save
    filename = 'manuel.docx'
    doc.save(filename)
    print(f"Manuel généré avec succès : {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_manual()
