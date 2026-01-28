"""
Word Export Module for ECDE Commercial Management
Handles generation of .docx reports mirroring existing PDF layouts.
"""

import os
from typing import Dict, Any, List

import shutil

from datetime import datetime
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx not installed.")
    Document = None

# ==================== HELPERS ====================

def check_logo_exists():
    "Check if logo exists, prioritizing logo_entete.png, then logo.png"
    # Helper duplicated from utils to avoid circular imports if utils imports this
    logo_names = ["logo_entete.png", "logo.png", "logo_gica.png"]
    for name in logo_names:
        if os.path.exists(name):
            return name
        # Check resource path manually if needed, but os.path.exists usually works
    return None

def format_currency(value):
    """Format currency (2 decimals, space separator) - Matches utils.format_currency"""
    if value is None: value = 0.0
    try:
        s = f"{float(value):,.2f}"
        return s.replace(",", " ").replace(".", ",")
    except:
        return "0,00"

def set_repeat_table_header(row):
    """Set table row as header to repeat on each page"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)
    return row

def set_cell_background(cell, hex_color):
    """Set cell background color (hex string without #)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_base_document(landscape=False):
    """Create a Document with A4 page size and standard margins"""
    if Document is None:
        return None
    
    doc = Document()
    section = doc.sections[0]
    
    # A4 Dimensions
    if landscape:
        section.page_height = Cm(21)
        section.page_width = Cm(29.7)
        section.orientation = WD_ALIGN_PARAGRAPH.LEFT # Actually enum values but this doesn't set orientation directly in python-docx properly without width/height swap
    else:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
    
    # Margins
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    
    return doc

def add_header(doc, title, subtitle=None, date_str=None, landscape=False):
    """Add standard header with Logo and Title"""
    # Use a table for layout: Logo Left, Text Center
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Widths
    logo_col_width = Cm(4)
    text_col_width = Cm(24 if landscape else 15)
    
    header_table.columns[0].width = logo_col_width
    header_table.columns[1].width = text_col_width
    
    # Logo
    logo_path = check_logo_exists()
    cell_logo = header_table.cell(0, 0)
    if logo_path:
        paragraph = cell_logo.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Cm(3.5), height=Cm(2.0))
        
    # Text
    cell_text = header_table.cell(0, 1)
    p = cell_text.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = p.add_run("GROUPE INDUSTRIEL DES CIMENTS D'ALGERIE\n")
    run1.font.name = 'Helvetica'
    run1.font.size = Pt(10)
    
    run2 = p.add_run("ENTREPRISE DES CIMENTS ET DERIVES D'ECH-CHELIFF\n")
    run2.font.name = 'Helvetica' # bold
    run2.bold = True
    run2.font.size = Pt(12)
    
    run3 = p.add_run(f"\n{title}\n")
    run3.font.name = 'Helvetica' # bold
    run3.bold = True
    run3.font.size = Pt(16)
    
    if subtitle:
        run4 = p.add_run(f"{subtitle}\n")
        run4.font.name = 'Helvetica'
        run4.font.size = Pt(11)
        
    if date_str:
        run5 = p.add_run(f"{date_str}")
        run5.font.name = 'Helvetica'
        run5.font.size = Pt(11)
        
    doc.add_paragraph() # Spacer

def style_table(table):
    """Apply standard borders and styling"""
    table.style = 'Table Grid'
    # Additional manual borders can be complicated in python-docx, relying on 'Table Grid' for now which puts borders everywhere.

def ensure_export_dir():
    directory = "Exports_Word"
    if not os.path.exists(directory):
        os.makedirs(directory)
    return directory

# ==================== REPORTS ====================

def generate_situation_word(data):
    """Situation Client"""
    doc = create_base_document(landscape=False)
    
    client = data['client']
    balance = data['balance']
    date_str = f"Date: {datetime.now().strftime('%d/%m/%Y')}"
    
    add_header(doc, "SITUATION CLIENT", date_str=date_str)
    
    # Client Info
    p = doc.add_paragraph()
    p.add_run(f"Client: {client['raison_sociale']}\n").bold = True
    p.add_run(f"Adresse: {client['adresse']}")
    
    doc.add_paragraph()
    
    # Balance Summary Table
    table = doc.add_table(rows=6, cols=2)
    style_table(table)
    table.autofit = False
    
    # Header
    row = table.rows[0]
    set_repeat_table_header(row)
    cell = row.cells[0]
    cell.merge(row.cells[1])
    cell.text = "Résumé Financier"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    set_cell_background(cell, "D3D3D3") # Light Grey
    
    # Data
    rows_data = [
        ("Report N-1:", balance['report']),
        ("Total Factures:", balance['total_factures']),
        ("Total Paiements:", balance['total_paiements']),
        ("Total Avoirs:", balance['total_avoirs']),
        ("SOLDE ACTUEL:", balance['solde'])
    ]
    
    for i, (label, val) in enumerate(rows_data):
        r = table.rows[i+1]
        r.cells[0].text = label
        r.cells[1].text = f"{format_currency(val)} DA" if i == 4 else format_currency(val)
        r.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if i == 4: # Solde Row
            r.cells[0].paragraphs[0].runs[0].bold = True
            r.cells[1].paragraphs[0].runs[0].bold = True
            set_cell_background(r.cells[0], "F5F5F5")
            set_cell_background(r.cells[1], "F5F5F5")

    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = "".join([c for c in client['raison_sociale'] if c.isalnum() or c in (' ', '_')]).strip()
    filename = os.path.join(directory, f"Situation_{clean_name}_{timestamp}.docx")
    
    doc.save(filename)
    return filename

def generate_daily_sales_word(data):
    """État Journalier des Ventes"""
    doc = create_base_document(landscape=True)
    
    date_fmt = data['date'] # Already formatted or YYYY-MM-DD? usually string
    # Move Header to Word Header Section for repetition
    section = doc.sections[0]
    header = section.header
    
    # Custom Header (Matched to Sales by Category style)
    header_table = header.add_table(rows=1, cols=3, width=Cm(27.7))
    header_table.autofit = False
    
    header_table.columns[0].width = Cm(3.0)
    header_table.columns[1].width = Cm(21.7) 
    header_table.columns[2].width = Cm(3.0)
    
    # Logo
    logo_path = check_logo_exists()
    cell_logo = header_table.cell(0, 0)
    if logo_path:
        p = cell_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2.9), height=Cm(1.8))
        
    # Text
    cell_text = header_table.cell(0, 1)
    p = cell_text.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = p.add_run("GROUPE INDUSTRIEL DES CIMENTS D'ALGERIE\n")
    run1.font.name = 'Cambria'
    run1.font.size = Pt(16) 
    run1.bold = True
    
    run2 = p.add_run("ENTREPRISE DES CIMENTS ET DERIVES D'ECH-CHELIFF\n")
    run2.font.name = 'Cambria'
    run2.bold = True
    run2.font.size = Pt(11) 
    
    # Title Section (Also in Header)
    p_title = header.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("\nETAT DES VENTES QUOTIDIENNES\n")
    r_title.font.name = 'Cambria'
    r_title.font.size = Pt(14)
    r_title.bold = True
    
    r_sub = p_title.add_run("Dépôt Oued Smar\n")
    r_sub.font.name = 'Cambria'
    r_sub.font.size = Pt(11)
    r_sub.bold = True
    r_sub.underline = True
    
    r_date = p_title.add_run(f"Journée du : {date_fmt}")
    r_date.font.name = 'Cambria'
    r_date.font.size = Pt(11)
    r_date.bold = True
    r_date.underline = True
    
    # Pagination in Footer
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.add_run("Page ")
    
    # Page field
    run_page = p_foot.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fldChar1)
    
    run_instr = p_foot.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run_instr._r.append(instrText)
    
    run_end = p_foot.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_end._r.append(fldChar2)
    
    p_foot.add_run(" / ")
    
    # NumPages field
    run_numpages = p_foot.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run_numpages._r.append(fldChar3)
    
    run_instr2 = p_foot.add_run()
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    run_instr2._r.append(instrText2)
    
    run_end2 = p_foot.add_run()
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run_end2._r.append(fldChar4)
    
    
    doc.add_paragraph() # Spacer in body
    
    # Table Details
    headers = ['Clients', 'Produit\nCode', 'Facture\nN.', 'Facture\nDate', 'Qte', 'HT', 'M.remise', 'TVA', 'TTC']
    # Widths approximation for A4 Landscape
    widths = [Cm(4.5), Cm(2.2), Cm(2.2), Cm(2.2), Cm(1.8), Cm(2.5), Cm(1.5), Cm(2.5), Cm(2.8)]
    
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    table.autofit = False
    
    # Header Row
    hdr_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Cambria'
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(hdr_cells[i], "D3D3D3")
        if i == 6: # M.remise Column (index 6 after removal)
             pass # Removed Yellow Background
        table.columns[i].width = widths[i]
        
    # Data
    for row in data['details']:
        tr = table.add_row()
        cells = tr.cells
        
        cells[0].text = str(row['client'])
        cells[1].text = str(row['code_produit'] or row['produit'][:10])
        cells[2].text = str(row['facture_num'])
        
        # Date fmt
        try:
            d_obj = datetime.strptime(row['date'], '%Y-%m-%d')
            d_txt = d_obj.strftime('%d-%b-%y')
        except:
            d_txt = row['date']
        cells[3].text = d_txt
        
        cells[4].text = f"{row['qte']:,.3f}".replace(",", " ").replace(".", ",")
        cells[5].text = format_currency(row['ht'])
        cells[6].text = "0,00"
        cells[7].text = format_currency(row['tva'])
        cells[8].text = format_currency(row['ttc'])
        
        # Font Cambria & MR Yellow
        for idx in range(9):
            p = cells[idx].paragraphs[0]
            if len(p.runs) > 0: p.runs[0].font.name = 'Cambria'
            else: p.add_run().font.name = 'Cambria' # Ensure run exists
            
            p.runs[0].font.size = Pt(10)
            p.runs[0].bold = True # Looks bold in screenshot
            
            if idx == 6: # M.remise
                pass # Removed Yellow


        # Align numbers
        # Align numbers
        for idx in range(4, 9):
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Total Row Details
    totals = data['totals']
    tr = table.add_row()
    cells = tr.cells
    cells[3].text = "TOTAL"
    cells[3].paragraphs[0].runs[0].bold = True
    cells[3].paragraphs[0].runs[0].font.name = 'Cambria'
    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    cells[4].text = f"{totals['day_qty']:,.3f}".replace(",", " ").replace(".", ",")
    cells[5].text = format_currency(totals['day_ht'])
    cells[6].text = "0,00"
    cells[7].text = format_currency(totals['day_tva'])
    cells[8].text = format_currency(totals['day_ttc'])
    
    for idx in range(4, 9):
         p = cells[idx].paragraphs[0]
         p.runs[0].bold = True
         p.runs[0].font.name = 'Cambria'
         p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
         if idx == 6: pass # set_cell_background(cells[idx], "FFFF00")
    
    
    doc.add_paragraph()
    
    # Summary Table (Products)
    headers_sum = ['Désignation', 'Qte Journée', 'Qte Cumulée']
    table_sum = doc.add_table(rows=1, cols=3)
    style_table(table_sum)
    table_sum.autofit = False
    
    set_repeat_table_header(table_sum.rows[0])
    h_cells = table_sum.rows[0].cells
    for i, h in enumerate(headers_sum):
        h_cells[i].text = h
        h_cells[i].paragraphs[0].runs[0].bold = True
        h_cells[i].paragraphs[0].runs[0].font.name = 'Cambria'
        set_cell_background(h_cells[i], "F5F5F5")
    
    # Increased by 15% from the 40% reduced size: 4.2 -> 4.83; 2.1 -> 2.415
    # Note: table.columns[].width is often ignored by Word if autofit is off but cells aren't set
    col_widths = [Cm(4.83), Cm(2.42), Cm(2.42)]
    for i in range(3):
        table_sum.columns[i].width = col_widths[i]
        h_cells[i].width = col_widths[i]

    for p in data['product_stats']:
        tr = table_sum.add_row()
        tr.cells[0].text = p['nom']
        tr.cells[1].text = f"{p['daily_qty']:,.3f}".replace(",", " ").replace(".", ",")
        tr.cells[2].text = f"{p['cumul_qty']:,.3f}".replace(",", " ").replace(".", ",")
        for i in range(3):
            tr.cells[i].width = col_widths[i]
            
    doc.add_paragraph()
    
    # Footer Totals
    # Footer Totals (Blue Bar)
    # Use a 2-col table to simulate the bar: Label Left, Value Right
    # "TOTAL ANNEE HT:"    "2 019 312,00 DA"
    
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.autofit = False
    footer_table.columns[0].width = Cm(15)
    footer_table.columns[1].width = Cm(10)
    
    fc = footer_table.rows[0].cells
    
    fc[0].text = "TOTAL ANNEE HT:"
    fc[1].text = f"{format_currency(totals['year_net_ht'])} DA"
    
    for c in fc:
        set_cell_background(c, "4F81BD") # Blue color from screenshot equivalent
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.name = 'Cambria'
        run.font.color.rgb = RGBColor(0, 0, 0) # Black text? Or White? Screenshot looks like Black text on Blue.
    
    fc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Ventes_Journalieres_{data['date']}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_sales_by_category_word(data, start_date, end_date):
    """État CA par Famille"""
    doc = create_base_document(landscape=True)
    
    # Custom Header for this report (Swapped sizes per user request)
    # Header Table: 3 columns to perfectly center the text [Logo, Text, Spacer]
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = False
    
    # Page width 29.7 - 2cm margins = 27.7cm
    # Adjusted to maximize center space while keeping logo fit:
    # Reduce side columns to 3.0cm to give 21.7cm to text.
    header_table.columns[0].width = Cm(3.0)
    header_table.columns[1].width = Cm(21.7) 
    header_table.columns[2].width = Cm(3.0)
    
    # Logo (Left)
    logo_path = check_logo_exists()
    cell_logo = header_table.cell(0, 0)
    if logo_path:
        paragraph = cell_logo.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run()
        # Slightly smaller logo to fit in 3cm col
        run.add_picture(logo_path, width=Cm(2.9), height=Cm(1.8))
        
    # Text (Center)
    cell_text = header_table.cell(0, 1)
    p = cell_text.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = p.add_run("GROUPE INDUSTRIEL DES CIMENTS D'ALGERIE\n")
    run1.font.name = 'Cambria'
    # First line Larger (16pt to match visual dominance)
    run1.font.size = Pt(16) 
    run1.bold = True
    
    run2 = p.add_run("ENTREPRISE DES CIMENTS ET DERIVES D'ECH-CHELIFF\n")
    run2.font.name = 'Cambria'
    run2.bold = True
    # Second line Smaller (11pt to ensure single line)
    run2.font.size = Pt(11) 
    
    run3 = p.add_run(f"\nÉTAT DU CHIFFRE D'AFFAIRES PÉRIODIQUE PAR FAMILLE D'ARTICLE\n")
    run3.font.name = 'Cambria'
    run3.bold = True
    run3.font.size = Pt(12) 
    
    run4 = p.add_run(f"Période: {start_date} au {end_date}")
    run4.font.name = 'Cambria'
    run4.bold = True # Added bold to match header style
    run4.font.size = Pt(11)
        
    doc.add_paragraph()
    
    grand_total = 0.0
    categories = ["Ciment"] + [k for k in data.keys() if k != "Ciment" and data[k]]
    
    for cat in categories:
        if cat not in data or not data[cat]: continue
        
        p = doc.add_paragraph(f"FAMILLE: {cat.upper()}")
        p.runs[0].bold = True
        p.runs[0].font.name = 'Cambria'
        p.runs[0].font.color.rgb = RGBColor(0x1a, 0x23, 0x7e) # Dark Blue
        
        # Table
        table = doc.add_table(rows=1, cols=3)
        style_table(table)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER # Center table
        
        # Headers: Designation, Qte, Montant HT
        headers = ["Désignation", "Quantité", "Montant HT"]
        
        row = table.rows[0]
        set_repeat_table_header(row)
        for i, text in enumerate(headers):
            row.cells[i].text = text
            row.cells[i].paragraphs[0].runs[0].bold = True
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(row.cells[i], "1A237E") # Dark Blue
            row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White
            
        table.columns[0].width = Cm(10.5) 
        table.columns[1].width = Cm(2.45)  
        table.columns[2].width = Cm(3.15)  
        
        subtotal = 0.0
        
        for idx, item in enumerate(data[cat]):
            tr = table.add_row()
            tr.cells[0].text = item['nom']
            tr.cells[0].paragraphs[0].runs[0].font.name = 'Cambria'
            
            tr.cells[1].text = f"{item['qte']:,.3f}".replace(",", " ").replace(".", ",")
            tr.cells[1].paragraphs[0].runs[0].font.name = 'Cambria'
            tr.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            tr.cells[2].text = format_currency(item['montant_ht'])
            tr.cells[2].paragraphs[0].runs[0].font.name = 'Cambria'
            tr.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Alternate Background Colors (White / Light Grey)
            if idx % 2 != 0: # Odd index = Even row number (2nd, 4th...) since 0 is 1st
                for c in tr.cells:
                    set_cell_background(c, "F5F5F5")
            
            subtotal += item['montant_ht']
            
        # Subtotal Row
        tr = table.add_row()
        tr.cells[1].text = "SOUS-TOTAL"
        tr.cells[1].paragraphs[0].runs[0].bold = True
        tr.cells[1].paragraphs[0].runs[0].font.name = 'Cambria'
        tr.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        tr.cells[2].text = format_currency(subtotal)
        tr.cells[2].paragraphs[0].runs[0].bold = True
        tr.cells[2].paragraphs[0].runs[0].font.name = 'Cambria'
        tr.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(tr.cells[2], "D3D3D3")
        
        grand_total += subtotal
        doc.add_paragraph()
        
    # Grand Total
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"TOTAL GÉNÉRAL HT: {format_currency(grand_total)} DA")
    run.bold = True
    run.font.name = 'Cambria'
    run.font.size = Pt(14)
    
    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(9.8)
    sig_table.columns[1].width = Cm(9.8)
    
    c1 = sig_table.cell(0, 0)
    p1 = c1.paragraphs[0]
    p1.add_run("LE FACTURIER").bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    c2 = sig_table.cell(0, 1)
    p2 = c2.paragraphs[0]
    p2.add_run("LE CHEF DE SERVICE COMMERCIAL").bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT # User said "a gauche", assuming left align in Right column?
    
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"CA_Par_Famille_{start_date}_{end_date}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_stock_valuation_word(data):
    """Rapport Stock Valorisé"""
    doc = create_base_document(landscape=True)
    
    prod_name = data['product']['nom']
    period_str = f"DU : {data['period']['start']}   AU : {data['period']['end']}"
    
    add_header(doc, "ETAT DES MOUVEMENTS DES STOCKS (VALORISES)", 
              subtitle=f"PRODUIT : {prod_name}\nUNITE : {data['product']['unite']}", 
              date_str=period_str, landscape=True)
              
    doc.add_paragraph()
    
    # Table structure
    # Header 1: JOURNEE | STOCK INITIAL | P.UNITAIRE | RECEPTIONS | VENTES | STOCK FINAL
    # Header 2:         | QTE | VAL     |            | QTE | VAL  | QTE | VAL | QTE | VAL
    
    # We will use a single header row for simplicity in Word or merge cells?
    # Word tables support merging.
    
    table = doc.add_table(rows=2, cols=10)
    style_table(table)
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])
    
    # Col Widths: 2.5cm each approx
    for i in range(10):
        table.columns[i].width = Cm(1.75)
        
    # Row 0
    r0 = table.rows[0]
    r0.cells[0].text = "JOURNEE"
    
    r0.cells[1].text = "STOCK INITIAL"
    r0.cells[1].merge(r0.cells[2])
    
    r0.cells[3].text = "P.UNITAIRE"
    # r0.cells[3].merge(r0.cells[3]) # No horizontal merge
    
    r0.cells[4].text = "RECEPTIONS"
    r0.cells[4].merge(r0.cells[5])
    
    r0.cells[6].text = "VENTES"
    r0.cells[6].merge(r0.cells[7])
    
    r0.cells[8].text = "STOCK FINAL"
    r0.cells[8].merge(r0.cells[9])
    
    # Vertical Merges (Journee, P.Unit)
    r0.cells[0].merge(table.rows[1].cells[0])
    r0.cells[3].merge(table.rows[1].cells[3])
    
    # Row 1 (Subheaders)
    r1 = table.rows[1]
    subheaders = ["", "QTE", "VAL", "", "QTE", "VAL", "QTE", "VAL", "QTE", "VAL"]
    for i, txt in enumerate(subheaders):
        if txt: r1.cells[i].text = txt
        
    # Style Headers
    for r in table.rows[:2]:
        for cell in r.cells:
            if cell.text.strip():
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                set_cell_background(cell, "D3D3D3")
                
    # Data
    for row in data['data']:
        tr = table.add_row()
        cells = tr.cells
        
        # Date
        try:
            d = datetime.strptime(row['date'], '%Y-%m-%d').strftime('%d/%m/%Y')
        except: d = row['date']
        cells[0].text = d
        
        cells[1].text = format_currency(row['stock_initial_qty'])
        cells[2].text = format_currency(row['stock_initial_val'])
        cells[3].text = format_currency(row['cout_achat'])
        cells[4].text = format_currency(row['reception_qty'])
        cells[5].text = format_currency(row['reception_val'])
        cells[6].text = format_currency(row['vente_qty'])
        cells[7].text = format_currency(row['vente_val'])
        cells[8].text = format_currency(row['stock_final_qty'])
        cells[9].text = format_currency(row['stock_final_val'])
        
        # Coloring Logic (Mirrors get_conditional_styles)
        # Pos > 0 Green, Neg < 0 Orange.
        # Word doesn't support text color easily without Oxml, stick to Black for now or use RGBColor.
        
        for i in range(1, 10):
            try:
                val = float(cells[i].text.replace(" ", "").replace(",", "."))
                p = cells[i].paragraphs[0]
                run = p.runs[0]
                if val > 0.001:
                    run.font.color.rgb = RGBColor(0, 128, 0) # Green
                elif val < -0.001:
                    run.font.color.rgb = RGBColor(255, 152, 0) # Orange
                else:
                    run.font.color.rgb = RGBColor(0, 0, 255) # Blue
            except: pass
            
            
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # TOTAL ROW
    t_recept_qty = sum([d['reception_qty'] for d in data['data']])
    t_recept_val = sum([d['reception_val'] for d in data['data']])
    t_vente_qty = sum([d['vente_qty'] for d in data['data']])
    t_vente_val = sum([d['vente_val'] for d in data['data']])
    
    tr_total = table.add_row()
    tr_total.cells[0].text = "TOTAL"
    tr_total.cells[0].paragraphs[0].runs[0].bold = True
    
    tr_total.cells[4].text = format_currency(t_recept_qty)
    tr_total.cells[5].text = format_currency(t_recept_val)
    tr_total.cells[6].text = format_currency(t_vente_qty)
    tr_total.cells[7].text = format_currency(t_vente_val)
    
    for c in tr_total.cells:
        set_cell_background(c, "E0E0E0")
        if c.text.strip():
             p = c.paragraphs[0]
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             if not p.runs: p.add_run(c.text)
             p.runs[0].bold = True
            
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_prod = "".join([c for c in prod_name if c.isalnum()]).strip()
    filename = os.path.join(directory, f"Stock_Valorise_{clean_prod}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_global_consumption_word(data, date_str):
    """État de Consommation Global"""
    doc = create_base_document(landscape=True)
    
    date_fmt = datetime.strptime(date_str, '%Y-%m-%d').strftime('%d/%m/%Y')
    add_header(doc, "ETAT DE CONSOMMATION GLOBAL", 
              subtitle=None, 
              date_str=f"A FIN : {date_fmt}", landscape=True)
              
    doc.add_paragraph()
    
    # Table Header:
    # Designation | U | JOURNEE | MOIS | ANNEE
    #             |   | Qte | Val | Qte | Val | Qte | Val
    
    table = doc.add_table(rows=2, cols=8)
    style_table(table)
    table.autofit = False
    
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])
    
    # Row 0
    r0 = table.rows[0]
    r0.cells[0].text = "Désignation"
    r0.cells[1].text = "U"
    r0.cells[2].text = "JOURNEE"
    r0.cells[2].merge(r0.cells[3])
    r0.cells[4].text = "CUMUL MOIS"
    r0.cells[4].merge(r0.cells[5])
    r0.cells[6].text = "CUMUL ANNEE"
    r0.cells[6].merge(r0.cells[7])
    
    # Vertical Merges
    r0.cells[0].merge(table.rows[1].cells[0])
    r0.cells[1].merge(table.rows[1].cells[1])
    
    # Row 1
    r1 = table.rows[1]
    subheaders = ["", "", "Qté", "Valeur", "Qté", "Valeur", "Qté", "Valeur"]
    for i, txt in enumerate(subheaders):
        if txt: r1.cells[i].text = txt
        
    # Styling Headers
    for r in table.rows[:2]:
        for c in r.cells:
            if c.text:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraphs[0].runs[0].bold = True
                set_cell_background(c, "D3D3D3")
                
    # Column Widths
    widths = [Cm(4.2), Cm(1.05), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
        
    for row in data['data']:
        tr = table.add_row()
        cells = tr.cells
        cells[0].text = row['product_name']
        cells[1].text = row['unit']
        
        cells[2].text = format_currency(row['daily_qty'])
        cells[3].text = format_currency(row['daily_val'])
        cells[4].text = format_currency(row['monthly_qty'])
        cells[5].text = format_currency(row['monthly_val'])
        cells[6].text = format_currency(row['yearly_qty'])
        cells[7].text = format_currency(row['yearly_val'])
        
        # Styling Numbers (2 to 7)
        for i in range(2, 8):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                val = float(cells[i].text.replace(" ", "").replace(",", "."))
                p = cells[i].paragraphs[0]
                run = p.runs[0]
                if val > 0.001:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif val < -0.001:
                    run.font.color.rgb = RGBColor(255, 152, 0)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 255)
            except: pass

    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Etat_Conso_Global_{date_str}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_movements_valorises_word(data_result, date_str):
    """Mouvements Valorisés"""
    data = data_result['data']
    totals = data_result['totals']
    
    doc = create_base_document(landscape=True)
    date_fmt = datetime.strptime(date_str, '%Y-%m-%d').strftime('%d/%m/%Y')
    add_header(doc, "ETAT DES MOUVEMENTS DES STOCKS VALORISES", 
              date_str=f"JOURNEE DU {date_fmt}", landscape=True)
    
    doc.add_paragraph()
    
    # Function to create table (Quantities or Values)
    def create_mv_table(title, is_value_table=False):
        doc.add_paragraph(title).runs[0].bold = True
        
        table = doc.add_table(rows=2, cols=12)
        style_table(table)
        table.autofit = False
        set_repeat_table_header(table.rows[0])
        set_repeat_table_header(table.rows[1])
        
        # Headers
        r0 = table.rows[0]
        r1 = table.rows[1]
        
        if is_value_table:
             h1 = ["Désignation", "Cout U.", "JOURNEE", "", "", "MOIS", "", "", "ANNEE", "", "", "VAL. FINALE"]
        else:
             h1 = ["Désignation", "U", "JOURNEE", "", "", "MOIS", "", "", "ANNEE", "", "", "STOCK FINAL"]
             
        # Merges
        r0.cells[0].text = h1[0]
        r0.cells[0].merge(r1.cells[0])
        
        r0.cells[1].text = h1[1]
        r0.cells[1].merge(r1.cells[1])
        
        # Journee
        r0.cells[2].text = "JOURNEE"
        r0.cells[2].merge(r0.cells[4]) # 2, 3, 4
        
        # Mois
        r0.cells[5].text = "MOIS"
        r0.cells[5].merge(r0.cells[7]) # 5, 6, 7
        
        # Annee
        r0.cells[8].text = "ANNEE"
        r0.cells[8].merge(r0.cells[10]) # 8, 9, 10
        
        # Final
        r0.cells[11].text = h1[11]
        r0.cells[11].merge(r1.cells[11])
        
        # Subheaders
        sub = ["", "", "S.Init", "Entrées", "Sorties"] * 3 # We need 3 sets? No.
        # Indices: 2,3,4 | 5,6,7 | 8,9,10
        for i in range(2, 11):
            lbl = ["S.Init", "Entrées", "Sorties"]
            r1.cells[i].text = lbl[(i-2)%3]
            
        # Style
        for r in table.rows[:2]:
            for c in r.cells:
                if c.text:
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].bold = True
                    set_cell_background(c, "D3D3D3")
                    
        # Widths
        # 12 cols, 28cm.
        # Des: 5cm, U: 2cm. Rem: 10 cols. 21cm left. ~2.1cm each.
        table.columns[0].width = Cm(3.5)
        table.columns[1].width = Cm(1.05)
        for i in range(2, 12):
            table.columns[i].width = Cm(1.5)
            
        return table
        
    # Table 1: Quantities
    t1 = create_mv_table("TABLEAU 1: QUANTITES")
    
    for row in data:
        tr = t1.add_row()
        tr.cells[0].text = row['designation']
        tr.cells[1].text = row['unite']
        tr.cells[2].text = format_currency(row['day']['init'])
        tr.cells[3].text = format_currency(row['day']['in'])
        tr.cells[4].text = format_currency(row['day']['out'])
        tr.cells[5].text = format_currency(row['month']['init'])
        tr.cells[6].text = format_currency(row['month']['in'])
        tr.cells[7].text = format_currency(row['month']['out'])
        tr.cells[8].text = format_currency(row['year']['init'])
        tr.cells[9].text = format_currency(row['year']['in'])
        tr.cells[10].text = format_currency(row['year']['out'])
        tr.cells[11].text = format_currency(row['final'])
        
        for k in range(2, 12):
            tr.cells[k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # Total T1
    tr = t1.add_row()
    tr.cells[0].text = "TOTAL"
    tr.cells[0].paragraphs[0].runs[0].bold = True
    tr.cells[2].text = format_currency(totals['day']['init'])
    tr.cells[3].text = format_currency(totals['day']['in'])
    tr.cells[4].text = format_currency(totals['day']['out'])
    tr.cells[5].text = format_currency(totals['month']['init'])
    tr.cells[6].text = format_currency(totals['month']['in'])
    tr.cells[7].text = format_currency(totals['month']['out'])
    tr.cells[8].text = format_currency(totals['year']['init'])
    tr.cells[9].text = format_currency(totals['year']['in'])
    tr.cells[10].text = format_currency(totals['year']['out'])
    tr.cells[11].text = format_currency(totals['final'])
    
    for c in tr.cells: set_cell_background(c, "E0E0E0")
    
    doc.add_paragraph()
    
    # Table 2: Values
    t2 = create_mv_table("TABLEAU 2: VALEURS (DA)", is_value_table=True)
    
    v_totals = [0.0] * 10
    
    for row in data:
        tr = t2.add_row()
        tr.cells[0].text = row['designation']
        tr.cells[1].text = format_currency(row['cout_unitaire'])
        tr.cells[2].text = format_currency(row['values']['day']['init'])
        tr.cells[3].text = format_currency(row['values']['day']['in'])
        tr.cells[4].text = format_currency(row['values']['day']['out'])
        tr.cells[5].text = format_currency(row['values']['month']['init'])
        tr.cells[6].text = format_currency(row['values']['month']['in'])
        tr.cells[7].text = format_currency(row['values']['month']['out'])
        tr.cells[8].text = format_currency(row['values']['year']['init'])
        tr.cells[9].text = format_currency(row['values']['year']['in'])
        tr.cells[10].text = format_currency(row['values']['year']['out'])
        tr.cells[11].text = format_currency(row['val_final'])
        
        for k in range(1, 12):
            tr.cells[k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Accumulate totals
        v_totals[0] += row['values']['day']['init']
        v_totals[1] += row['values']['day']['in']
        v_totals[2] += row['values']['day']['out']
        v_totals[3] += row['values']['month']['init']
        v_totals[4] += row['values']['month']['in']
        v_totals[5] += row['values']['month']['out']
        v_totals[6] += row['values']['year']['init']
        v_totals[7] += row['values']['year']['in']
        v_totals[8] += row['values']['year']['out']
        v_totals[9] += row['val_final']

    # Total T2
    tr = t2.add_row()
    tr.cells[0].text = "TOTAL"
    tr.cells[0].paragraphs[0].runs[0].bold = True
    
    # Fill totals
    for i, val in enumerate(v_totals):
        tr.cells[i+2].text = format_currency(val) # +2 because index 0 is Des, 1 is Cout U (skipped for total?)
    
    for c in tr.cells: set_cell_background(c, "E0E0E0")

    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Etat_Mouvements_Stocks_Valorises_{date_str}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_annual_receivables_word(data, date_str):
    """État Annuel des Créances"""
    doc = create_base_document(landscape=True)
    date_fmt = datetime.strptime(date_str, "%Y-%m-%d").strftime("%d/%m/%Y")
    
    add_header(doc, "ÉTAT RÉCAPITULATIF ANNUEL DES CRÉANCES ET RECOUVREMENTS CLIENTS", 
              date_str=f"SITUATION AU {date_fmt}", landscape=True)
              
    doc.add_paragraph()
    
    headers = ["Raison Sociale", "Solde au 01/01", "Achats (Année)", "Paiements (Année)", "Solde Final", "% Recouvrement"]
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    table.autofit = False
    
    set_repeat_table_header(table.rows[0])
    
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(table.rows[0].cells[i], "D3D3D3")
        
    # Widths 
    # [5, 4, 4, 4, 4, 3] = 24cm -- plenty of room
    widths = [Cm(4.2), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.1)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
        
    for row in data['data']:
        tr = table.add_row()
        tr.cells[0].text = row['raison_sociale']
        tr.cells[1].text = format_currency(row['solde_01_01'])
        tr.cells[2].text = format_currency(row['achats'])
        tr.cells[3].text = format_currency(row['paiements'])
        tr.cells[4].text = format_currency(row['solde_final'])
        tr.cells[5].text = f"{row['recouvrement']:.1f}%"
        
        for k in range(1, 6):
            tr.cells[k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
    # Totals
    totals = data['totals']
    tr = table.add_row()
    tr.cells[0].text = f"SOLDE GLOBAL AU {date_fmt}"
    tr.cells[0].paragraphs[0].runs[0].bold = True
    
    tr.cells[1].text = format_currency(totals['solde_init'])
    tr.cells[2].text = format_currency(totals['achats'])
    tr.cells[3].text = format_currency(totals['paiements'])
    tr.cells[4].text = format_currency(totals['solde_final'])
    tr.cells[5].text = ""
    
    for c in tr.cells: 
        set_cell_background(c, "E0E0E0")
        c.paragraphs[0].runs[0].bold = True
        
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Etat_Creances_Annuelles_{date_str}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_cancellations_analysis_word(annulations):
    """Analyse des Annulations (Specific Visuals)"""
    doc = create_base_document(landscape=False)
    
    add_header(doc, "ANALYSE DES FACTURES ANNULÉES")
    doc.add_paragraph()
    
    if not annulations:
        doc.add_paragraph("Aucune annulation enregistrée.")
    else:
        # Table of cancellations
        headers = ["DATE", "NUMERO", "MONTANT HT", "MOTIF"]
        table = doc.add_table(rows=1, cols=4)
        style_table(table)
        table.autofit = False
        set_repeat_table_header(table.rows[0])
        
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_background(table.rows[0].cells[i], "D3D3D3")
            
        table.columns[0].width = Cm(2.1)
        table.columns[1].width = Cm(2.1)
        table.columns[2].width = Cm(2.8)
        table.columns[3].width = Cm(5.6)
        
        total_perdu = 0.0
        by_motif = {}
        
        for row in annulations:
             # Normalize data (tuple vs dict)
            try:
                date_val = str(row['date_annulation'])[:10]
                num = str(row['numero_facture'])
                motif = row['motif']
                mht = row['montant_original_ht']
            except:
                date_val = str(row[2])[:10]
                num = str(row[5])
                motif = row[4]
                mht = row[6]
            
            total_perdu += mht
            if motif not in by_motif: by_motif[motif] = 0
            by_motif[motif] += 1
            
            tr = table.add_row()
            cells = tr.cells
            cells[0].text = date_val
            cells[1].text = num
            cells[2].text = format_currency(mht)
            cells[3].text = motif
            
            # Styling: Orange & Bold
            for c in cells:
                p = c.paragraphs[0]
                run = p.runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 140, 0) # Dark Orange
            
            # Add watermark text "ANNULÉE" in cells? No, per spec: "Si l'état concerne une facture annulée..."
            # For this report, the whole report is about cancellations.
            # "Le document Word doit simplement refléter l'affichage visuel existant (Police Orange et Grasse)" -> Done above.
            
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"TOTAL MANQUE À GAGNER (HT): {format_currency(total_perdu)} DA")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 0, 0) # Red
        
        doc.add_paragraph()
        doc.add_paragraph("SYNTHÈSE PAR MOTIF:").runs[0].bold = True
        
        for m, count in by_motif.items():
            doc.add_paragraph(f"- {m}: {count} fois", style='List Bullet')

    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Analyse_Annulations_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_grand_livre_word(data, period):
    """
    Grand Livre Détaillé des Opérations Clients.
    One section per client.
    Table: DATE | REF | LIBELLE | DEBIT | CREDIT | SOLDE
    """
    doc = create_base_document(landscape=True) # Landscape for more width
    
    start_date = period['start']
    end_date = period['end']
    
    add_header(doc, "GRAND-LIVRE DÉTAILLÉ DES OPÉRATIONS CLIENTS", 
              date_str=f"SITUATION DU {datetime.strptime(start_date, '%Y-%m-%d').strftime('%d/%m/%Y')} AU {datetime.strptime(end_date, '%Y-%m-%d').strftime('%d/%m/%Y')}", 
              landscape=True)
    
    doc.add_paragraph()
    
    # Check if empty
    if not data:
        doc.add_paragraph("Aucune donnée trouvée pour la période sélectionnée.")
    
    for client_data in data:
        client = client_data['client']
        
        # Header for Client
        # Using a Paragraph with Shading? Or simple Bold text
        p = doc.add_paragraph()
        run = p.add_run(f"CLIENT: {client['raison_sociale']} ({client['code_client'] or 'N/A'})")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        
        # Table
        # Cols: Date(2.5), Ref(2.5), Libelle(10), Debit(3), Credit(3), Solde(3) ~ 24cm
        headers = ["DATE", "RÉF", "LIBELLÉ DE L'ÉCRITURE", "DÉBIT", "CRÉDIT", "SOLDE"]
        
        table = doc.add_table(rows=1, cols=6)
        style_table(table)
        table.autofit = False
        
        # Header Row
        hr = table.rows[0]
        set_repeat_table_header(hr)
        
        for i, h in enumerate(headers):
            hr.cells[i].text = h
            hr.cells[i].paragraphs[0].runs[0].bold = True
            hr.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(hr.cells[i], "D3D3D3")
            
        # Widths
        widths = [Cm(2.5), Cm(2.5), Cm(10.0), Cm(3.0), Cm(3.0), Cm(3.0)]
        for i, w in enumerate(widths):
            table.columns[i].width = w
            
        # 1. Initial Balance Row
        r_init = table.add_row()
        r_init.cells[0].text = datetime.strptime(start_date, '%Y-%m-%d').strftime('%d/%m/%Y')
        r_init.cells[2].text = "SOLDE INITIAL AU DÉBUT DE PÉRIODE"
        r_init.cells[2].paragraphs[0].runs[0].bold = True
        
        init_bal = client_data['initial_balance']
        
        # If Initial Balance 
        # < 0 (Debt) -> Black (Default)
        # > 0 (Advance) -> Blue
        r_init.cells[5].text = format_currency(init_bal)
        if init_bal > 0.001: 
             r_init.cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255) # Blue
        else:
             r_init.cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0) # Black
        
        set_cell_background(r_init.cells[2], "F0F0F0")
        
        # 2. Movements
        for mv in client_data['movements']:
            tr = table.add_row()
            cells = tr.cells
            
            # Date
            try:
                d_fmt = datetime.strptime(mv['date'], '%Y-%m-%d').strftime('%d/%m/%Y')
            except: d_fmt = mv['date']
            cells[0].text = d_fmt
            
            # Ref
            cells[1].text = str(mv['ref'])
            
            # Libelle
            cells[2].text = mv['libelle']
            
            # Debit / Credit
            if mv['debit'] != 0:
                cells[3].text = format_currency(mv['debit'])
            else:
                cells[3].text = "-"
                
            if mv['credit'] != 0:
                cells[4].text = format_currency(mv['credit'])
            else:
                cells[4].text = "-"
                
            # Solde Progressif
            cells[5].text = format_currency(mv['solde_progressif'])
            
            # Styling
            # Align Numbers
            for idx in [3, 4, 5]:
                 cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Formatting Cancelled
            if mv.get('is_cancelled'):
                # Orange & Bold for specific cells or whole row?
                # "Les factures Annulées doivent rester en Orange... avec affichage du Motif"
                for c in cells:
                    p = c.paragraphs[0]
                    if len(p.runs) == 0: p.add_run(c.text) # Refresh runs
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 140, 0) # Orange
                        run.bold = True
            else:
                # Color Solde
                try:
                    val = float(mv['solde_progressif'])
                    if val > 0.001: # Advance
                        cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255) # Blue
                    else:
                        cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0) # Black
                except: pass

        # 3. Total Period Row
        r_tot = table.add_row()
        r_tot.cells[2].text = "TOTAUX PÉRIODE"
        r_tot.cells[2].paragraphs[0].runs[0].bold = True
        r_tot.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        r_tot.cells[3].text = format_currency(client_data['total_debit'])
        r_tot.cells[4].text = format_currency(client_data['total_credit'])
        
        r_tot.cells[3].paragraphs[0].runs[0].bold = True
        r_tot.cells[4].paragraphs[0].runs[0].bold = True
        r_tot.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_tot.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        set_cell_background(r_tot.cells[3], "E0E0E0")
        set_cell_background(r_tot.cells[4], "E0E0E0")
        
        # 4. Final Balance Row
        r_fin = table.add_row()
        r_fin.cells[2].text = "SOLDE FINAL AU " + datetime.strptime(end_date, '%Y-%m-%d').strftime('%d/%m/%Y')
        r_fin.cells[2].paragraphs[0].runs[0].bold = True
        r_fin.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        fin_bal = client_data['final_balance']
        r_fin.cells[5].text = format_currency(fin_bal)
        if len(r_fin.cells[5].paragraphs[0].runs) > 0:
             r_fin.cells[5].paragraphs[0].runs[0].bold = True
             if fin_bal > 0.001: # Advance
                 r_fin.cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255) # Blue
             else:
                 r_fin.cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0) # Black
        
        set_cell_background(r_fin.cells[5], "FFE599") # Light Orange/Yellow highlight
             
        doc.add_page_break()
        
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Grand_Livre_{start_date}_{end_date}_{timestamp}.docx")
    doc.save(filename)
    return filename

def generate_recovery_word(data, month, year):
    """
    Generate Word for Suivi Recouvrement Mensuel (M-1)
    """
    if not data or not data.get('data'):
        return None
        
    doc = create_base_document(landscape=False)
    
    add_header(doc, f"ÉTAT DE COUVERTURE DES CRÉANCES", 
               subtitle=f"(MOIS {month:02d}/{year})", 
               date_str=f"Date: {datetime.now().strftime('%d/%m/%Y')}", landscape=False)
               
    doc.add_paragraph()
    
    # Global Stats
    if data.get('totals'):
        rate = data['totals'].get('rate', 0.0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Taux de Recouvrement Global : {rate:.2f}%")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        
        # Color Logic
        if rate >= 80:
            run.font.color.rgb = RGBColor(0, 128, 0) # Green
        elif rate > 50:
            run.font.color.rgb = RGBColor(255, 152, 0) # Orange
        else:
            run.font.color.rgb = RGBColor(255, 0, 0) # Red
            
    doc.add_paragraph()
    
    # Table
    headers = ["Client", "Dette M-1\n(Cible)", "Paiements M\n(Réalisé)", "Reste à Payer\n(Ecart)", "Statut"]
    
    table = doc.add_table(rows=1, cols=5)
    style_table(table)
    table.autofit = False
    
    # Header Row
    row = table.rows[0]
    set_repeat_table_header(row)
    
    col_widths = [Cm(6), Cm(3.25), Cm(3.25), Cm(3.25), Cm(3.25)] # Total ~19cm
    
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        set_cell_background(cell, "D3D3D3")
        table.columns[i].width = col_widths[i]
        
    # Data
    for row_data in data['data']:
        tr = table.add_row()
        cells = tr.cells
        
        cells[0].text = row_data['raison_sociale']
        cells[1].text = format_currency(row_data['dette_m_1'])
        cells[2].text = format_currency(row_data['paiements_m'])
        cells[3].text = format_currency(row_data['reste_a_payer'])
        cells[4].text = row_data['statut']
        
        # Alignments
        for i in range(1, 4):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Coloring Logic
        statut = row_data['statut']
        color = None
        bold = False
        
        if statut == "RÉGLÉ":
            color = RGBColor(0, 128, 0) # Green
        elif statut == "ALERTE RECOUVREMENT":
            color = RGBColor(255, 0, 0) # Red
            bold = True
        else:
            color = RGBColor(255, 152, 0) # Orange
            
        if color:
            # Apply to Client Name and Status
            for idx in [0, 4]: 
                p = cells[idx].paragraphs[0]
                if not p.runs: p.add_run(cells[idx].text)
                for r in p.runs:
                    r.font.color.rgb = color
                    if bold: r.bold = True
                        
    # Totals Row
    if data.get('totals'):
        t = data['totals']
        tr = table.add_row()
        cells = tr.cells
        cells[0].text = "TOTAL"
        cells[1].text = format_currency(t['target'])
        cells[2].text = format_currency(t['realized'])
        
        for c in cells:
            set_cell_background(c, "E0E0E0")
            p = c.paragraphs[0]
            if not p.runs: p.add_run(c.text)
            p.runs[0].bold = True
            
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.autofit = False
    
    # 6cm each
    sig_table.columns[0].width = Cm(6)
    sig_table.columns[1].width = Cm(7)
    sig_table.columns[2].width = Cm(6)
    
    sig_table.cell(0, 0).text = "Service Recouvrement"
    sig_table.cell(0, 2).text = "Direction Commerciale"
    
    for c in sig_table.rows[0].cells:
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not p.runs: p.add_run(c.text)
        p.runs[0].bold = True
        
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Etat_Recouvrement_{month}-{year}_{timestamp}.docx")
    

    doc.save(filename)
    return filename

def generate_pareto_word(data: Dict[str, Any], start_date: str, end_date: str):
    """
    Generate Pareto Analysis Word Document.
    """
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from utils import check_logo_exists, format_currency, generate_pareto_charts, ensure_export_dir
    
    doc = Document()
    
    # Header
    HEADER_TEXT = "ANALYSE DE PERFORMANCE COMMERCIALE (PARETO)"
    
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Cm(18))
    header_table.autofit = False
    header_table.columns[0].width = Cm(4)
    header_table.columns[1].width = Cm(10)
    header_table.columns[2].width = Cm(4)
    
    # Logo
    logo_path = check_logo_exists()
    if logo_path:
        cell_logo = header_table.cell(0, 0)
        p = cell_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
        
    # Title
    cell_title = header_table.cell(0, 1)
    p_title = cell_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(HEADER_TEXT)
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.name = 'Cambria'
    run_title.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    
    # Date Range
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"Période du : {start_date} au {end_date}")
    run_sub.font.size = Pt(11)
    
    doc.add_paragraph() # Spacer
    
    # Charts
    curve_path, pie_path = generate_pareto_charts(data['data'])
    
    if curve_path and os.path.exists(curve_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(curve_path, width=Inches(6))
        
        # Pareto Legend
        p_leg = doc.add_paragraph()
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_leg_title = p_leg.add_run("Interprétation du Diagramme :\n")
        run_leg_title.bold = True
        run_leg_title.font.size = Pt(9)
        
        run_leg_text = p_leg.add_run("Les barres classent vos clients par volume de chiffre d'affaires. La ligne rouge montre le cumul. La zone où la ligne coupe la barre des 80% identifie vos clients stratégiques (Classe A).")
        run_leg_text.font.size = Pt(9)
        
    doc.add_paragraph()
    
    if pie_path and os.path.exists(pie_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(pie_path, width=Inches(4))
        
        # ABC Legend
        p_abc = doc.add_paragraph()
        p_abc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # A
        run_a_title = p_abc.add_run("Classe A (Vert) : ")
        run_a_title.bold = True
        run_a_title.font.size = Pt(9)
        run_a_title.font.color.rgb = RGBColor(46, 125, 50)
        
        run_a_text = p_abc.add_run("Représente 80% de votre activité. Ce sont vos clients piliers. Toute baisse de leur part est un risque majeur pour l'unité.\n")
        run_a_text.font.size = Pt(9)
        
        # B
        run_b_title = p_abc.add_run("Classe B (Bleu) : ")
        run_b_title.bold = True
        run_b_title.font.size = Pt(9)
        run_b_title.font.color.rgb = RGBColor(21, 101, 192)
        
        run_b_text = p_abc.add_run("Représente les 15% suivants. Ce sont vos clients en développement ou à fort potentiel.\n")
        run_b_text.font.size = Pt(9)
        
        # C
        run_c_title = p_abc.add_run("Classe C (Gris) : ")
        run_c_title.bold = True
        run_c_title.font.size = Pt(9)
        run_c_title.font.color.rgb = RGBColor(117, 117, 117)
        
        run_c_text = p_abc.add_run("Représente les derniers 5%. Ce sont des clients occasionnels qui génèrent un grand volume administratif pour un faible revenu.")
        run_c_text.font.size = Pt(9)
        
    doc.add_paragraph()
    
    # Dynamic Synthesis
    clients_a = 0
    total_clients = len(data['data'])
    for row in data['data']:
         if row['classe'] == 'A': clients_a += 1
         
    if total_clients > 0:
        perc_clients_a = (clients_a / total_clients) * 100
        
        # Use a single-cell table for shading
        synth_table = doc.add_table(rows=1, cols=1)
        synth_table.autofit = False
        synth_table.columns[0].width = Cm(17)
        synth_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cell = synth_table.cell(0, 0)
        set_cell_background(cell, "E8EAF6") # Light Indigo background
        
        p_synth = cell.paragraphs[0]
        p_synth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_synth_title = p_synth.add_run("Analyse de l'unité : ")
        run_synth_title.bold = True
        run_synth_title.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
        
        run_synth_text = p_synth.add_run(f"Votre activité est concentrée sur {clients_a} clients qui réalisent à eux seuls 80% du chiffre d'affaires.")
        run_synth_text.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
        
        if clients_a < 5 or perc_clients_a < 10:
             run_warn = p_synth.add_run(" Attention à la forte dépendance envers ce petit groupe.")
             run_warn.bold = True
             run_warn.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
             
        doc.add_paragraph() # Spacer

    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Column widths
    widths = [Cm(1.5), Cm(8), Cm(4), Cm(2.5), Cm(2)]
    headers = ['Rang', 'Client', 'Chiffre d\'Affaires', '% Cumulé', 'Classe']
    
    hdr_cells = table.rows[0].cells
    for i, (text, width) in enumerate(zip(headers, widths)):
        hdr_cells[i].text = text
        hdr_cells[i].width = width
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "1A237E") # Dark Blue
        
    # Data
    for row in data['data']:
        
        cells = table.add_row().cells
        cells[0].text = str(row['rank'])
        cells[1].text = row['client_name']
        cells[2].text = format_currency(row['ca'])
        cells[3].text = f"{row['cumul_perc']:.2f}%"
        cells[4].text = row['classe']
        
        # Alignments
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Colors
        bg_color = "FFFFFF"
        text_color_rgb = RGBColor(0, 0, 0)
        
        if row['classe'] == 'A':
            bg_color = "E8F5E9" # Light Green
            text_color_rgb = RGBColor(46, 125, 50) # Dark Green
        elif row['classe'] == 'B':
            bg_color = "E3F2FD" # Light Blue
            text_color_rgb = RGBColor(21, 101, 192) # Dark Blue
            
        for i in range(5):
            set_cell_background(cells[i], bg_color)
            
        # Set Class Text Color
        if row['classe'] in ['A', 'B']:
            if len(cells[4].paragraphs[0].runs) > 0:
                run = cells[4].paragraphs[0].runs[0]
            else:
                run = cells[4].paragraphs[0].add_run(row['classe'])
            run.font.bold = True
            run.font.color.rgb = text_color_rgb

    # Table is enough for summary logic, no need for footer summary text anymore since we have the detailed synthesis above.
    
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(directory, f"Pareto_Word_{start_date}_{end_date}_{timestamp}.docx")
    doc.save(filename)
    
    return filename

def generate_etat_104_word(data: list, start_date: str, end_date: str):
    """
    Générer document Word pour État 104 (Ventes par client).
    
    Format miroir du PDF:
    - En-tête avec logo GICA
    - Titre : "ETAT 104 DE L'ANNÉE : [année]"
    - Tableau avec totaux
    """
    from datetime import datetime
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt, RGBColor, Inches
    
    if Document is None:
        raise ImportError("python-docx non installé")
    
    doc = create_base_document(landscape=True)
    
    # En-tête avec logo et titre
    year = datetime.strptime(start_date, "%Y-%m-%d").year
    add_header(doc, f"ÉTAT 104 DE L'ANNÉE : {year}", 
               subtitle=f"Période : {start_date} au {end_date}", 
               landscape=True)
    
    # Tableau de données
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    style_table(table)
    
    # En-têtes
    headers = table.rows[0].cells
    header_texts = [
        "N° d'ordre",
        "Raison Sociale / Nom du Client",
        "Adresse précise",
        "NIF (15 chiffres)",
        "Article d'imposition (A)",
        "Montant des ventes (HT)",
        "Montant de la TVA",
        "Montant TTC"
    ]
    
    for i, text in enumerate(header_texts):
        headers[i].text = text
        headers[i].paragraphs[0].runs[0].bold = True
        headers[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        headers[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(headers[i], '1a237e')
    
    # Données et calcul totaux
    total_ht = 0.0
    total_tva = 0.0
    total_ttc = 0.0
    
    for row_data in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_data['numero'])
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        row_cells[1].text = row_data['raison_sociale']
        
        row_cells[2].text = row_data['adresse']
        
        row_cells[3].text = row_data['nif']
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        row_cells[4].text = row_data['article_imposition']
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        row_cells[5].text = format_currency(row_data['total_ht'])
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        row_cells[6].text = format_currency(row_data['total_tva'])
        row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        row_cells[7].text = format_currency(row_data['total_ttc'])
        row_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        total_ht += row_data['total_ht']
        total_tva += row_data['total_tva']
        total_ttc += row_data['total_ttc']
    
    # Ligne de total
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = ""
    total_row[2].text = ""
    total_row[3].text = ""
    total_row[4].text = "TOTAL"
    total_row[4].paragraphs[0].runs[0].bold = True
    total_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    total_row[5].text = format_currency(total_ht)
    total_row[5].paragraphs[0].runs[0].bold = True
    total_row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    total_row[6].text = format_currency(total_tva)
    total_row[6].paragraphs[0].runs[0].bold = True
    total_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    total_row[7].text = format_currency(total_ttc)
    total_row[7].paragraphs[0].runs[0].bold = True
    total_row[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Mettre en surbrillance la ligne de total
    for cell in total_row:
        set_cell_background(cell, 'e3f2fd')
    
    # Répéter l'en-tête si multi-pages
    set_repeat_table_header(table.rows[0])
    
    # Ajuster largeurs des colonnes
    table.columns[0].width = Inches(0.6)   # N°
    table.columns[1].width = Inches(2.5)   # Raison Sociale
    table.columns[2].width = Inches(2.2)   # Adresse
    table.columns[3].width = Inches(1.2)   # NIF
    table.columns[4].width = Inches(1.0)   # Article
    table.columns[5].width = Inches(1.3)   # HT
    table.columns[6].width = Inches(1.3)   # TVA
    table.columns[7].width = Inches(1.3)   # TTC
    
    # Sauvegarder
    export_dir = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(export_dir, f"Etat_104_{year}_{timestamp}.docx")
    doc.save(filename)
    
    print(f"✅ Word État 104 généré : {filename}")
    return filename


def generate_cockpit_word(data: Dict[str, Any]):

    """
    Generate Master Dashboard (Cockpit) Word Document.
    """
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from utils import check_logo_exists, format_currency, generate_cockpit_charts, ensure_export_dir
    
    doc = Document()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Setup Page margins for Landscape-ish feel (though easier to keep Portrait for max compatibility, 
    # user asked for "Réplique fidèle", usually A4 Landscape is better for Dashboards)
    section = doc.sections[0]
    section.orientation = 1 # Landscape
    # Accessing underlying XML to truly force Landscape if needed, but python-docx SECTION orientation is simpler
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    
    # Header
    # ------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(4)
    header_table.columns[1].width = Cm(20)
    
    logo_path = check_logo_exists()
    if logo_path:
        cell_logo = header_table.cell(0, 0)
        p = cell_logo.paragraphs[0]
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(3.5))
        
    cell_title = header_table.cell(0, 1)
    p_title = cell_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TABLEAU DE BORD MAÎTRE (COCKPIT)")
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    
    p_sub = cell_title.add_paragraph(f"Période: {data.get('period', '--')}")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # 1. Tiles (Table with Backgrounds)
    # ---------------------------------
    # 1 Row, 4 Cols
    t_tiles = doc.add_table(rows=1, cols=4)
    t_tiles.autofit = False
    col_width = Cm(6.5)
    
    kpis = data['kpis']
    evo = kpis['evolution']
    arrow = "▲" if evo >= 0 else "▼"
    
    tiles_info = [
        ("PERFORMANCE VENTES", format_currency(kpis['ca_curr']), f"{arrow} {abs(evo):.1f}% vs M-1", "1E88E5"),
        ("SANTÉ FINANCIÈRE", f"{kpis['recovery_rate']:.1f}%", "Recouvrement", "43A047"),
        ("RISQUE CRÉANCE (+30J)", format_currency(kpis['debt_30_days']), "Montant à Risque", "FB8C00"),
        ("ALERTE OPÉRATIONNELLE", f"{kpis['cancel_rate']:.1f}%", "Taux Annulation", "E53935")
    ]
    
    row_cells = t_tiles.rows[0].cells
    for i, (title, val, sub, color) in enumerate(tiles_info):
        cell = row_cells[i]
        cell.width = col_width
        set_cell_background(cell, color)
        
        # Title
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(title)
        r1.font.color.rgb = RGBColor(255, 255, 255)
        r1.font.bold = True
        r1.font.size = Pt(9)
        
        # Value
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(val)
        r2.font.color.rgb = RGBColor(255, 255, 255)
        r2.font.bold = True
        r2.font.size = Pt(18)
        
        # Sub
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r3 = p3.add_run(sub)
        r3.font.color.rgb = RGBColor(255, 255, 255)
        r3.font.size = Pt(9)
        
    doc.add_paragraph()
    
    # 2. Charts
    # ---------
    from utils import generate_cockpit_charts
    p_a, p_b, p_c = generate_cockpit_charts(data)
    
    t_charts = doc.add_table(rows=1, cols=3)
    t_charts.autofit = False
    
    if p_a and os.path.exists(p_a):
        c = t_charts.cell(0, 0)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(p_a, width=Cm(5))
        
    if p_b and os.path.exists(p_b):
        c = t_charts.cell(0, 1)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(p_b, width=Cm(8))
        
    if p_c and os.path.exists(p_c):
        c = t_charts.cell(0, 2)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(p_c, width=Cm(11))
        
    doc.add_paragraph()
    
    # 3. Alerts
    # ---------
    p_alert = doc.add_paragraph()
    run_alert = p_alert.add_run("ALERTES DE SÉCURITÉ (Code Rouge)")
    run_alert.bold = True
    run_alert.font.color.rgb = RGBColor(255, 0, 0)
    run_alert.font.size = Pt(12)
    
    if data['alerts']:
        t_alerts = doc.add_table(rows=1, cols=3)
        t_alerts.style = 'Table Grid'
        t_alerts.autofit = False
        t_alerts.columns[0].width = Cm(12)
        t_alerts.columns[1].width = Cm(6)
        t_alerts.columns[2].width = Cm(8)
        
        hdrs = t_alerts.rows[0].cells
        hdrs[0].text = "Client"
        hdrs[1].text = "Montant à Risque"
        hdrs[2].text = "Motif"
        
        for h in hdrs:
            set_cell_background(h, "FFEBEE")
            p = h.paragraphs[0]
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 0, 0)
                
        for alert in data['alerts']:
            row = t_alerts.add_row().cells
            row[0].text = alert['name']
            row[1].text = format_currency(alert['amount'])
            row[2].text = alert['reason']
            
            # Style rows
            for cell in row:
                set_cell_background(cell, "FFEBEE") # Red background for all alert rows
                p = cell.paragraphs[0]
                if len(p.runs) > 0:
                    p.runs[0].font.color.rgb = RGBColor(0, 0, 0) # Black text
                else:
                    # Accessing underlying run creation if strict needed
                    pass
    else:
        doc.add_paragraph("Aucune alerte critique détectée.")
        
    directory = ensure_export_dir()
    filename = os.path.join(directory, f"Cockpit_Word_{timestamp}.docx")
    doc.save(filename)
    return filename


def generate_invoice_word(invoice_data):
    """
    Générer Facture au format Word (OVERLAY MODE - Données Uniquement)
    Pour impression sur papier pré-imprimé.
    """
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from utils import nombre_en_lettres
    
    if Document is None:
        return None
        
    doc = create_base_document(landscape=True)
    
    # --- HEADER SECTION (Top Right Data) ---
    # We need to position "N° Facture" and "Date" at specific coordinates.
    # Using a borderless table is the safest way in standard Word flow.
    
    # Grid: [Empty Space Left] [Invoice Info Right]
    # Approx 20cm empty, then Info
    
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(20) # Spacer to push right
    header_table.columns[1].width = Cm(7)  # Info area
    
    cell_info = header_table.cell(0, 1)
    p = cell_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Align left within the right box
    
    # Add vertical spacing to push down to the pre-printed box line?
    # "tu doit imprimer ou générer uniquement les informations... a l'endroit exact"
    # Assuming standard header height from margin is sufficient, or adding Breaks.
    # Let's add simple data.
    
    # Invoice Number
    run = p.add_run(f"\n\n{invoice_data['numero']}\n") # Extra newlines for vertical alignment
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Facture Num (Official)
    run = p.add_run(f"\n{invoice_data.get('numero_facture', invoice_data['numero'])}\n")
    run.font.name = 'Cambria'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Date
    if invoice_data.get('date_facture'):
        try:
             dt = datetime.strptime(invoice_data['date_facture'], '%Y-%m-%d')
             date_str = dt.strftime('%d/%m/%Y')
        except:
             date_str = invoice_data['date_facture']
    else:
        date_str = ""
    
    run = p.add_run(f"\n{date_str}")
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    
    doc.add_paragraph() # Spacer
    
    # --- CLIENT & PAYMENT INFO (Middle Section) ---
    # Two Columns.
    # Left: Client Data (No labels like "Client:", just the value "SARL FOO")
    # Right: Payment Data
    
    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Cm(14)
    info_table.columns[1].width = Cm(13)
    
    # Left Cell: Client
    c1 = info_table.cell(0, 0)
    p = c1.paragraphs[0]
    # Add spacing to align with pre-printed lines
    
    def add_val(paragraph, value, space_before=0):
        if space_before:
            paragraph.add_run("\n" * space_before)
        r = paragraph.add_run(f"   {value}\n") # Indent slightly
        r.font.name = 'Cambria'
        r.font.size = Pt(11)
        r.bold = True
        
    # Vertical spacing matching the Lines in the photo approx
    # Line 1: Raison Sociale
    add_val(p, invoice_data['raison_sociale']) 
    
    # Line 2: Adresse
    add_val(p, invoice_data['adresse'])
    
    # Line 3: RC / NIS / NIF / ART (Often on same line or 2 lines)
    # Combining them linearly
    combined_ids = f"{invoice_data['rc']}                  {invoice_data['nis']}                  {invoice_data['nif']}                  {invoice_data.get('article_imposition', '')}"
    add_val(p, combined_ids)
    
    # Right Cell: Payment (Mode, Banque...)
    c2 = info_table.cell(0, 1)
    p = c2.paragraphs[0]
    
    # Skip lines to align with "Mode de règlement" usually lower down
    p.add_run("\n\n") 
    add_val(p, "Virement") # Mode
    add_val(p, "BNA OUED SMAR") # Banque
    add_val(p, "001 00634 0300 000 519 61") # Compte
    
    doc.add_paragraph()
    
    # --- PRODUCT TABLE (Data Only) ---
    # 5 Cols. No Headers. Widths must match exactly.
    widths = [Cm(11), Cm(2), Cm(3), Cm(4), Cm(5)]
    
    prod_table = doc.add_table(rows=1, cols=5)
    prod_table.autofit = False
    
    # No Header Row created or populated. Directly data.
    
    for ligne in invoice_data['lignes']:
        row = prod_table.add_row()
        cells = row.cells
        
        # Set Widths
        for i, w in enumerate(widths):
            cells[i].width = w
            
        cells[0].text = ligne['product_nom']
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        cells[1].text = ligne['unite']
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format numbers with spaces
        cells[2].text = f"{ligne['quantite']:,.3f}".replace(",", " ").replace(".", ",")
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        pu = ligne.get('prix_unitaire', 0.0)
        cells[3].text = format_currency(pu)
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        cells[4].text = format_currency(ligne['montant'])
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Styling
        for cell in cells:
            p = cell.paragraphs[0]
            if p.runs:
                run = p.runs[0]
            else:
                run = p.add_run(cell.text)
            run.font.name = 'Cambria'
            run.font.size = Pt(11)
            run.bold = True
            
        # Optional Remise line handled as separate row or text?
        if hasattr(ligne, 'get') and ligne.get('taux_remise', 0) > 0:
             p_rem = cells[0].add_paragraph(f"Remise {ligne['taux_remise']:.0f}%")
             p_rem.alignment = WD_ALIGN_PARAGRAPH.CENTER
             p_rem.runs[0].font.name = 'Cambria'
             p_rem.runs[0].font.size = Pt(10)

    # --- TOTALS SECTION ---
    # Need to align Bottom-Right
    doc.add_paragraph()
    doc.add_paragraph() # Spacer for table bottom
    
    footer_table = doc.add_table(rows=1, cols=2)
    footer_table.autofit = False
    footer_table.columns[0].width = Cm(18)
    footer_table.columns[1].width = Cm(9)
    
    # Left: Driver Info and Signature First, then Amount in Words Below
    left_cell = footer_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    
    # Driver Info (Values Only) - MOVED TO TOP
    if invoice_data.get('chauffeur'):
        p.add_run(f"   {invoice_data['chauffeur']}\n").font.name = 'Cambria'
    if invoice_data.get('matricule'):
        p.add_run(f"   {invoice_data['matricule']}\n").font.name = 'Cambria'
        
    # Signer Name - MOVED BEFORE AMOUNT  
    p.add_run("\n\n   HAMMICHE MAKHLOUF\n\n")
    
    # Amount in Words (No "Arrêtée..." text) - MOVED TO BOTTOM
    montant_ttc = invoice_data['montant_ttc']
    run = p.add_run(f"\n{nombre_en_lettres(montant_ttc)}\n")
    run.bold = True
    run.italic = True
    run.font.name = 'Cambria'
    run.font.size = Pt(12) 
    
    # Right: Totals Box (Values Only)
    right_cell = footer_table.cell(0, 1)
    
    # Nested table for alignment
    t_totals = right_cell.add_table(rows=3, cols=1)
    # Height spacing is critical here.
    
    def add_total(idx, val, bold=False):
        c = t_totals.rows[idx].cells[0]
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(format_currency(val))
        r.font.name = 'Cambria'
        r.font.size = Pt(12)
        r.bold = bold
        
    add_total(0, invoice_data['montant_ht'])
    add_total(1, invoice_data['montant_tva'])
    add_total(2, invoice_data['montant_ttc'], bold=True)
    
    directory = ensure_export_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = "".join([c for c in invoice_data['raison_sociale'] if c.isalnum() or c in (' ', '_')]).strip()
    filename = os.path.join(directory, f"Facture_Overlay_{invoice_data['numero']}_{clean_name}_{timestamp}.docx")
    
    doc.save(filename)
    return filename
